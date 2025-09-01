# Export Manager module for Enhanced RAG Chatbot
import os
import re
import tempfile
from datetime import datetime

# Optional imports for export functionality
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("ReportLab not available - PDF export will use text fallback")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not available - Word export will use text fallback")

from config import SUPPORTED_EXPORT_FORMATS

class ExportManager:
    """Handles exporting conversations, answers, and summaries to various formats"""
    
    def __init__(self):
        # Determine available formats based on installed packages
        self.supported_formats = ["txt"]  # Always available
        
        if REPORTLAB_AVAILABLE:
            self.supported_formats.append("pdf")
        
        if DOCX_AVAILABLE:
            self.supported_formats.append("docx")
        
        self.temp_dir = tempfile.mkdtemp()
        print(f"Export formats available: {self.supported_formats}")
    
    def export_conversation(self, conversation_history, format_type="pdf", filename=None):
        """Export conversation history to specified format"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"conversation_{timestamp}"
        
        try:
            if format_type == "pdf":
                if REPORTLAB_AVAILABLE:
                    filepath = self._export_to_pdf(conversation_history, filename)
                    return (filepath, f"Conversation exported to PDF: {os.path.basename(filepath)}")
                else:
                    return (None, "PDF export not available. ReportLab package not installed. Falling back to TXT format.")
            elif format_type == "docx":
                if DOCX_AVAILABLE:
                    filepath = self._export_to_docx(conversation_history, filename)
                    return (filepath, f"Conversation exported to Word: {os.path.basename(filepath)}")
                else:
                    return (None, "Word export not available. python-docx package not installed. Falling back to TXT format.")
            elif format_type == "txt":
                filepath = self._export_to_txt(conversation_history, filename)
                return (filepath, f"Conversation exported to text: {os.path.basename(filepath)}")
            else:
                raise ValueError(f"Unsupported format: {format_type}")
        except Exception as e:
            return (None, f"Export failed: {str(e)}")
    
    def export_summary(self, document_summary, format_type="pdf", filename=None):
        """Export document summary to specified format"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"summary_{timestamp}"
        
        try:
            if format_type == "pdf":
                if REPORTLAB_AVAILABLE:
                    filepath = self._export_summary_to_pdf(document_summary, filename)
                    return (filepath, f"Summary exported to PDF: {os.path.basename(filepath)}")
                else:
                    return (None, "PDF export not available. ReportLab package not installed. Falling back to TXT format.")
            elif format_type == "docx":
                if DOCX_AVAILABLE:
                    filepath = self._export_summary_to_docx(document_summary, filename)
                    return (filepath, f"Summary exported to Word: {os.path.basename(filepath)}")
                else:
                    return (None, "Word export not available. python-docx package not installed. Falling back to TXT format.")
            elif format_type == "txt":
                filepath = self._export_summary_to_txt(document_summary, filename)
                return (filepath, f"Summary exported to text: {os.path.basename(filepath)}")
            else:
                raise ValueError(f"Unsupported format: {format_type}")
        except Exception as e:
            return (None, f"Export failed: {str(e)}")
    
    def _export_to_pdf(self, conversation, filename):
        """Export conversation to PDF"""
        filepath = os.path.join(self.temp_dir, f"{filename}.pdf")
        
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # Center
            textColor='#1a1a1a'
        )
        story.append(Paragraph("Document Q&A Conversation", title_style))
        story.append(Spacer(1, 25))
        
        # Add timestamp
        timestamp_style = ParagraphStyle(
            'Timestamp',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=20,
            alignment=1,  # Center
            textColor='#666666'
        )
        timestamp = datetime.now().strftime("Generated on %B %d, %Y at %I:%M %p")
        story.append(Paragraph(timestamp, timestamp_style))
        story.append(Spacer(1, 30))
        
        # Conversation content
        for i, (question, answer) in enumerate(conversation, 1):
            # Question header
            q_header_style = ParagraphStyle(
                'QuestionHeader',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=8,
                spaceBefore=15,
                textColor='#2E86AB',
                fontName='Helvetica-Bold'
            )
            story.append(Paragraph(f"Question {i}", q_header_style))
            
            # Question text
            q_style = ParagraphStyle(
                'Question',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12,
                leftIndent=0,
                textColor='#333333',
                fontName='Helvetica'
            )
            story.append(Paragraph(question, q_style))
            
            # Answer header
            a_header_style = ParagraphStyle(
                'AnswerHeader',
                parent=styles['Heading3'],
                fontSize=12,
                spaceAfter=8,
                spaceBefore=15,
                textColor='#28a745',
                fontName='Helvetica-Bold'
            )
            story.append(Paragraph("Answer:", a_header_style))
            
            # Answer text with better formatting
            a_style = ParagraphStyle(
                'Answer',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=25,
                leftIndent=20,
                textColor='#1a1a1a',
                fontName='Helvetica',
                leading=14,  # Line spacing
                alignment=0  # Left align
            )
            
            # Clean and format the answer text
            formatted_answer = self._format_answer_text(answer)
            story.append(Paragraph(formatted_answer, a_style))
            
            # Add separator line between Q&A pairs
            if i < len(conversation):
                story.append(Spacer(1, 10))
                story.append(Paragraph("<hr/>", styles['Normal']))
                story.append(Spacer(1, 10))
        
        doc.build(story)
        return filepath
    
    def _format_answer_text(self, text):
        """Format answer text for better readability"""
        if not text:
            return ""
        
        # Remove debug information and metadata
        text = re.sub(r'\*\*Method Used:\*\*.*?\)', '', text, flags=re.DOTALL)
        text = re.sub(r'\*\*Confidence Score:\*\*.*?\)', '', text, flags=re.DOTALL)
        text = re.sub(r'\*\*Response Time:\*\*.*?\)', '', text, flags=re.DOTALL)
        text = re.sub(r'\*\*Retrieved Chunks:\*\*.*?\)', '', text, flags=re.DOTALL)
        
        # Clean up extra whitespace and line breaks
        text = re.sub(r'\n+', '\n', text)
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'---\s*', '', text)
        
        # Ensure proper sentence endings
        text = text.strip()
        if text and not text.endswith(('.', '!', '?')):
            text += '.'
        
        return text
    
    def _export_to_docx(self, conversation, filename):
        """Export conversation to Word document"""
        filepath = os.path.join(self.temp_dir, f"{filename}.docx")
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Document Q&A Conversation', 0)
        title.alignment = 1  # Center alignment
        
        # Add timestamp
        timestamp = doc.add_paragraph()
        timestamp.alignment = 1  # Center alignment
        timestamp.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}").italic = True
        
        doc.add_paragraph()  # Add spacing
        
        # Conversation content
        for i, (question, answer) in enumerate(conversation, 1):
            # Question section
            q_heading = doc.add_heading(f'Question {i}', level=1)
            q_heading.style.font.color.rgb = RGBColor(46, 134, 171)  # Blue color
            
            # Question text
            q_text = doc.add_paragraph(question)
            q_text.style.font.size = Pt(11)
            q_text.style.font.name = 'Calibri'
            
            # Answer section
            a_heading = doc.add_heading('Answer:', level=2)
            a_heading.style.font.color.rgb = RGBColor(40, 167, 69)  # Green color
            
            # Answer text with formatting
            formatted_answer = self._format_answer_text(answer)
            a_text = doc.add_paragraph(formatted_answer)
            a_text.style.font.size = Pt(10)
            a_text.style.font.name = 'Calibri'
            a_text.paragraph_format.left_indent = Inches(0.2)
            
            # Add spacing between Q&A pairs
            if i < len(conversation):
                doc.add_paragraph()  # Add spacing
                # Add a simple separator line
                separator = doc.add_paragraph('_' * 50)
                separator.alignment = 1  # Center alignment
                doc.add_paragraph()  # Add spacing
        
        doc.save(filepath)
        return filepath
    
    def _export_to_txt(self, conversation, filename):
        """Export conversation to text file"""
        filepath = os.path.join(self.temp_dir, f"{filename}.txt")
        
        with open(filepath, 'w', encoding='utf-8') as f:
            # Header
            f.write("Document Q&A Conversation\n")
            f.write("=" * 50 + "\n")
            f.write(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n")
            f.write("=" * 50 + "\n\n")
            
            # Conversation content
            for i, (question, answer) in enumerate(conversation, 1):
                f.write(f"Question {i}\n")
                f.write("-" * 20 + "\n")
                f.write(f"{question}\n\n")
                
                f.write("Answer:\n")
                f.write("-" * 20 + "\n")
                formatted_answer = self._format_answer_text(answer)
                f.write(f"{formatted_answer}\n\n")
                
                # Add separator between Q&A pairs
                if i < len(conversation):
                    f.write("_" * 50 + "\n\n")
        
        return filepath
    
    def _export_summary_to_pdf(self, summary, filename):
        """Export document summary to PDF"""
        filepath = os.path.join(self.temp_dir, f"{filename}.pdf")
        
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'SummaryTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1
        )
        story.append(Paragraph("Document Summary", title_style))
        story.append(Spacer(1, 20))
        
        # Summary content
        summary_style = ParagraphStyle(
            'Summary',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12
        )
        
        # Split summary into paragraphs
        paragraphs = summary.split('\n\n')
        for para in paragraphs:
            if para.strip():
                story.append(Paragraph(para.strip(), summary_style))
        
        doc.build(story)
        return filepath
    
    def _export_summary_to_docx(self, summary, filename):
        """Export document summary to Word document"""
        filepath = os.path.join(self.temp_dir, f"{filename}.docx")
        
        doc = Document()
        doc.add_heading('Document Summary', 0)
        
        # Split summary into paragraphs
        paragraphs = summary.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        doc.save(filepath)
        return filepath
    
    def _export_summary_to_txt(self, summary, filename):
        """Export document summary to text file"""
        filepath = os.path.join(self.temp_dir, f"{filename}.txt")
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("Document Summary\n")
            f.write("=" * 30 + "\n\n")
            f.write(summary)
        
        return filepath
    
    def cleanup_temp_files(self):
        """Clean up temporary files"""
        try:
            import shutil
            shutil.rmtree(self.temp_dir, ignore_errors=True)
            self.temp_dir = tempfile.mkdtemp()
        except Exception as e:
            print(f"Cleanup failed: {e}")
    
    def get_export_formats(self):
        """Get list of supported export formats"""
        return self.supported_formats
